# -*- coding: utf-8 -*-
"""Небольшая обёртка над python-docx под оформление исполнительной документации."""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"


def new_doc(landscape=False, margins=(1.5, 1.0, 1.5, 1.0)):
    """margins = (top, right, bottom, left) в см."""
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.element.rPr.rFonts.set(qn("w:cs"), FONT)
    pf = st.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0

    s = doc.sections[0]
    if landscape:
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = Cm(29.7), Cm(21.0)
    else:
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.top_margin, s.right_margin = Cm(margins[0]), Cm(margins[1])
    s.bottom_margin, s.left_margin = Cm(margins[2]), Cm(margins[3])
    return doc


ALIGN = {
    "l": WD_ALIGN_PARAGRAPH.LEFT,
    "c": WD_ALIGN_PARAGRAPH.CENTER,
    "r": WD_ALIGN_PARAGRAPH.RIGHT,
    "j": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def p(doc, text="", size=11, bold=False, italic=False, align="j",
      before=0, after=3, indent=None, caps=False, underline=False):
    par = doc.add_paragraph()
    par.alignment = ALIGN[align]
    par.paragraph_format.space_before = Pt(before)
    par.paragraph_format.space_after = Pt(after)
    if indent is not None:
        par.paragraph_format.first_line_indent = Cm(indent)
    if text:
        r = par.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.underline = underline
        r.font.all_caps = caps
        r.font.name = FONT
        r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return par


def hint(doc, text, size=8):
    """Пояснительная подпись под строкой формы (курсивом, по центру)."""
    return p(doc, f"({text})", size=size, italic=True, align="c", after=6)


def _set_bottom_border(par):
    ppr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    ppr.append(pbdr)


def line(doc, text="", size=11, align="l", bold=False, after=0, min_empty=True):
    """Строка формы: значение с подчёркиванием по всей ширине."""
    par = p(doc, text if text else (" " if min_empty else ""),
            size=size, align=align, bold=bold, after=after)
    _set_bottom_border(par)
    return par


def field(doc, label, value="", note=None, label_size=10, value_size=11, lines=1):
    """Поле формы: подпись → строка(и) со значением → курсивная подсказка."""
    if label:
        p(doc, label, size=label_size, align="l", after=1)
    line(doc, value, size=value_size)
    for _ in range(lines - 1):
        line(doc, "")
    if note:
        hint(doc, note)
    else:
        p(doc, "", after=4)


def blank_lines(doc, n=1, size=11):
    for _ in range(n):
        line(doc, "", size=size)


def title(doc, text, size=13):
    p(doc, text, size=size, bold=True, align="c", after=4)


def table(doc, headers, rows, widths=None, size=9, header_size=9,
          align="c", first_col_align="c", body_align="l", row_height=None):
    """row_height — минимальная высота строки в см (для бланков под заполнение)."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        par = hdr[i].paragraphs[0]
        par.alignment = ALIGN["c"]
        par.paragraph_format.space_after = Pt(0)
        r = par.add_run(h)
        r.bold = True
        r.font.size = Pt(header_size)
        r.font.name = FONT
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            par = cells[i].paragraphs[0]
            par.alignment = ALIGN[first_col_align if i == 0 else body_align]
            par.paragraph_format.space_after = Pt(0)
            r = par.add_run(str(val))
            r.font.size = Pt(size)
            r.font.name = FONT
    if row_height:
        from docx.enum.table import WD_ROW_HEIGHT_RULE
        for row in t.rows[1:]:
            row.height = Cm(row_height)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    if widths:
        # ширины растягиваются на всю полезную ширину полосы набора
        sec = doc.sections[-1]
        usable = Emu(sec.page_width - sec.left_margin - sec.right_margin).cm
        total = sum(widths)
        if total > 0:
            widths = [w * usable / total for w in widths]
        # фиксированная раскладка — иначе Word/LibreOffice игнорируют ширины колонок
        t.autofit = False
        tbl_pr = t._tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tbl_pr.append(layout)
        for i, w in enumerate(widths):
            t.columns[i].width = Cm(w)
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def signature_block(doc, roles, size=10):
    """roles: список подписей вида (должность/роль, пояснение)."""
    for role, note in roles:
        p(doc, role, size=size, align="l", after=2)
        t = doc.add_table(rows=1, cols=3)
        t.autofit = False
        cells = t.rows[0].cells
        cells[0].width, cells[1].width, cells[2].width = Cm(6.0), Cm(4.5), Cm(6.5)
        for c, txt in zip(cells, ["", "", ""]):
            c.text = txt
        for i, sub in enumerate(["должность", "подпись", "расшифровка подписи"]):
            par = cells[i].paragraphs[0]
            par.alignment = ALIGN["c"]
            par.paragraph_format.space_after = Pt(0)
            _set_bottom_border(par)
            par.add_run(" ")
            sp = cells[i].add_paragraph()
            sp.alignment = ALIGN["c"]
            sp.paragraph_format.space_after = Pt(0)
            sr = sp.add_run(f"({sub})")
            sr.italic = True
            sr.font.size = Pt(7)
            sr.font.name = FONT
        if note:
            hint(doc, note, size=7.5)
        else:
            p(doc, "", after=6)


def page_break(doc):
    doc.add_page_break()


def new_section(doc, landscape=False, margins=(1.5, 1.0, 1.5, 1.0)):
    """Новый раздел с заданной ориентацией — для сводных файлов на печать."""
    from docx.enum.section import WD_SECTION
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    if landscape:
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = Cm(29.7), Cm(21.0)
    else:
        s.orientation = WD_ORIENT.PORTRAIT
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.top_margin, s.right_margin = Cm(margins[0]), Cm(margins[1])
    s.bottom_margin, s.left_margin = Cm(margins[2]), Cm(margins[3])
    return s


def footer_note(doc, text):
    p(doc, text, size=8, italic=True, align="l", before=8, after=0)
