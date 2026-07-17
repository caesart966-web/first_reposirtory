"""Чтение документов: PDF, DOCX, XLSX и разбор графика.

Каждый парсер возвращает структуру с текстом/таблицами и никогда не бросает на
битом файле — вместо этого поле `error`. read_schedule/parse_estimate — эвристики
для типовых таблиц; результаты идут на проверку человеку (шаг 3 workflow).
"""
from __future__ import annotations

from pathlib import Path


def read_pdf(path: str) -> dict:
    try:
        import pdfplumber
    except ImportError:  # pragma: no cover
        return {"error": "pdfplumber не установлен", "text": "", "pages": []}
    if not Path(path).exists():
        return {"error": f"file not found: {path}", "text": "", "pages": []}
    pages = []
    try:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                pages.append({"page": i + 1, "text": page.extract_text() or ""})
    except Exception as exc:  # pragma: no cover
        return {"error": f"pdf parse error: {exc}", "text": "", "pages": []}
    return {"text": "\n".join(p["text"] for p in pages), "pages": pages, "error": None}


def read_docx(path: str) -> dict:
    try:
        import docx
    except ImportError:  # pragma: no cover
        return {"error": "python-docx не установлен", "text": "", "paragraphs": [], "tables": []}
    if not Path(path).exists():
        return {"error": f"file not found: {path}", "text": "", "paragraphs": [], "tables": []}
    try:
        d = docx.Document(path)
    except Exception as exc:  # pragma: no cover
        return {"error": f"docx parse error: {exc}", "text": "", "paragraphs": [], "tables": []}
    paragraphs = [p.text for p in d.paragraphs]
    tables = [[[c.text for c in row.cells] for row in t.rows] for t in d.tables]
    return {"text": "\n".join(paragraphs), "paragraphs": paragraphs, "tables": tables, "error": None}


def read_xlsx(path: str, sheet: str | None = None) -> dict:
    try:
        import openpyxl
    except ImportError:  # pragma: no cover
        return {"error": "openpyxl не установлен", "sheets": []}
    if not Path(path).exists():
        return {"error": f"file not found: {path}", "sheets": []}
    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    except Exception as exc:  # pragma: no cover
        return {"error": f"xlsx parse error: {exc}", "sheets": []}
    sheets = []
    for ws in wb.worksheets:
        if sheet and ws.title != sheet:
            continue
        rows = [["" if v is None else str(v) for v in row]
                for row in ws.iter_rows(values_only=True)]
        sheets.append({"name": ws.title, "rows": rows})
    wb.close()
    return {"sheets": sheets, "error": None}


# ---- эвристический разбор графика ----

_START_KEYS = ("начал", "старт", "start", "нач.")
_FINISH_KEYS = ("оконч", "финиш", "finish", "end", "кон.")
_NAME_KEYS = ("наимен", "работа", "задача", "task", "name", "вид работ")
_RESP_KEYS = ("ответствен", "подрядчик", "responsible", "исполнитель")


def _find_header(rows: list[list[str]]) -> tuple[int, dict]:
    for idx, row in enumerate(rows[:15]):
        low = [c.lower() for c in row]
        cols = {}
        for j, cell in enumerate(low):
            if any(k in cell for k in _NAME_KEYS) and "name" not in cols:
                cols["name"] = j
            elif any(k in cell for k in _START_KEYS) and "start" not in cols:
                cols["start"] = j
            elif any(k in cell for k in _FINISH_KEYS) and "finish" not in cols:
                cols["finish"] = j
            elif any(k in cell for k in _RESP_KEYS) and "responsible" not in cols:
                cols["responsible"] = j
        if "name" in cols and ("start" in cols or "finish" in cols):
            return idx, cols
    return -1, {}


def _norm_date(v: str) -> str | None:
    v = (v or "").strip()
    if not v:
        return None
    # ISO уже?
    if len(v) >= 10 and v[4] == "-" and v[7] == "-":
        return v[:10]
    for sep in (".", "/"):
        if sep in v:
            parts = v.split(" ")[0].split(sep)
            if len(parts) == 3:
                d, m, y = parts
                if len(y) == 4:
                    return f"{y}-{int(m):02d}-{int(d):02d}"
    return None


def read_schedule(path: str) -> dict:
    """Разбирает XLSX-график в список частичных ScheduleTask (без id). Эвристика по заголовкам."""
    parsed = read_xlsx(path)
    if parsed.get("error"):
        return {"error": parsed["error"], "tasks": []}
    for sheet in parsed["sheets"]:
        rows = sheet["rows"]
        hidx, cols = _find_header(rows)
        if hidx < 0:
            continue
        tasks = []
        for row in rows[hidx + 1:]:
            if len(row) <= cols["name"] or not row[cols["name"]].strip():
                continue
            task = {"name": row[cols["name"]].strip()}
            if "start" in cols and len(row) > cols["start"]:
                task["planned_start"] = _norm_date(row[cols["start"]])
            if "finish" in cols and len(row) > cols["finish"]:
                task["planned_finish"] = _norm_date(row[cols["finish"]])
            if "responsible" in cols and len(row) > cols["responsible"]:
                task["responsible"] = row[cols["responsible"]].strip() or None
            tasks.append(task)
        return {"tasks": tasks, "sheet": sheet["name"], "error": None,
                "note": "эвристический разбор; проверьте на шаге 3 workflow"}
    return {"tasks": [], "error": "не найден заголовок графика (наименование/сроки)"}
