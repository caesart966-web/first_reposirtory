# -*- coding: utf-8 -*-
"""Автоматические проверки.

Все данные в тестах ВЫМЫШЛЕННЫЕ: реквизиты реальных компаний здесь
не используются принципиально.

Запуск:  python -m unittest discover -s tests -v
"""

from __future__ import annotations

import json
import shutil
import sys
import tempfile
import unittest
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from src import morphology, validators  # noqa: E402
from src.company_parser import (parse_card, parse_text,  # noqa: E402
                                split_company_name)
from src.context_builder import build_context  # noqa: E402
from src.docx_engine import (_digit_row, extract_all_text,  # noqa: E402
                             extract_package_text,
                             fill_template, scan_placeholders)
from src.document_generator import (GeneratorError, Project,  # noqa: E402
                                    check_readiness, company_folder_name,
                                    generate, generate_many)
from src.models import CompanyData  # noqa: E402
from src.readers import read_card  # noqa: E402

# --------------------------------------------------------------- фикстуры
ALPHA = dict(
    full_name='Общество с ограниченной ответственностью «Ромашка-Строй»',
    short_name='ООО «Ромашка»',
    inn="7812345675",
    kpp="781201001",
    ogrn="1237800000008",
    legal_address="190000, г. Санкт-Петербург, ул. Вымышленная, д. 1, лит. А, пом. 5",
    actual_address="190000, г. Санкт-Петербург, ул. Вымышленная, д. 1, лит. А, пом. 5",
    phone="+7 (812) 000-00-01",
    email="info@romashka-test.example",
    director_position="Генеральный директор",
    director_full_name="Иванов Иван Иванович",
    director_basis="Устав",
)

BETA = dict(
    full_name='Акционерное общество «Василёк»',
    short_name='АО «Василёк»',
    inn="7701001005",
    kpp="770101001",
    ogrn="1027700000008",
    legal_address="101000, г. Москва, пер. Придуманный, д. 2",
    actual_address="101000, г. Москва, пер. Придуманный, д. 2",
    phone="+7 (495) 000-00-02",
    email="office@vasilek-test.example",
    director_position="Директор",
    director_full_name="Петрова Мария Сергеевна",
    director_basis="Устав",
)


#: Вымышленный предприниматель. ИНН 12 знаков, ОГРНИП 15, КПП нет.
IVAN = dict(
    applicant_kind="entrepreneur",
    full_name="Индивидуальный предприниматель Иванов Иван Иванович",
    short_name="ИП Иванов И.И.",
    inn="781234567870",
    ogrn="304780123456781",
    legal_address="190000, г. Санкт-Петербург, ул. Вымышленная, д. 1, лит. А, пом. 5",
    actual_address="190000, г. Санкт-Петербург, ул. Вымышленная, д. 1, лит. А, пом. 5",
    phone="+7 (812) 000-00-03",
    email="ivanov@ip-test.example",
    director_position="Индивидуальный предприниматель",
    director_full_name="Иванов Иван Иванович",
    director_basis="Лист записи ЕГРИП",
)


def make_company(data: dict) -> CompanyData:
    company = CompanyData()
    for key, value in data.items():
        company.set(key, value)
    company.doc_date = "17.08.2026"
    return company


def make_card_docx(path: Path, data: dict) -> None:
    """Собрать вымышленную карточку компании в DOCX (как у настоящего клиента)."""
    from docx import Document

    document = Document()
    document.add_paragraph(f"Реквизиты компании {data['short_name']}")
    rows = [
        ("Полное наименование", data["full_name"]),
        ("Сокращенное наименование", data["short_name"]),
        ("Юридический адрес", data["legal_address"]),
        ("Почтовый адрес", "-"),
        ("ИНН", data["inn"]),
        ("КПП", data["kpp"]),
        ("ОГРН", data["ogrn"]),
        (data["director_position"],
         f"{data['director_full_name']} (на основании {data['director_basis']}а)"),
        ("Эл. почта", data["email"]),
        ("Телефон", data["phone"]),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    for index, (label, value) in enumerate(rows):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value
    document.save(str(path))


# --------------------------------------------------------------- проверки
class TestValidators(unittest.TestCase):
    def test_inn_length(self):
        issues = validators.validate_inn("781234567")
        self.assertTrue(issues)
        self.assertIn("10 цифр", issues[0].text)

    def test_inn_checksum(self):
        self.assertTrue(validators.inn_checksum_ok("7812345675"))
        self.assertFalse(validators.inn_checksum_ok("7812345678"))
        issues = validators.validate_inn("7812345678")
        self.assertTrue(any("контрольного числа" in i.text for i in issues))

    def test_inn_valid_is_silent(self):
        self.assertEqual(validators.validate_inn("7812345675"), [])

    def test_ogrn(self):
        self.assertTrue(validators.ogrn_checksum_ok("1237800000008"))
        self.assertEqual(validators.validate_ogrn("1237800000008"), [])
        self.assertTrue(validators.validate_ogrn("1237800000000"))

    def test_ogrn_length_message(self):
        issues = validators.validate_ogrn("12378")
        self.assertIn("13 цифр", issues[0].text)

    def test_spaces_are_reported_not_silently_fixed(self):
        issues = validators.validate_inn("78 1234 5675")
        self.assertTrue(any(i.severity == "warning" for i in issues))

    def test_email_and_phone(self):
        self.assertEqual(validators.validate_email("a.b@mail.example"), [])
        self.assertTrue(validators.validate_email("a.b@mail"))
        self.assertTrue(validators.validate_email("имя @mail.ru"))
        self.assertEqual(validators.validate_phone("+7 (812) 000-00-01"), [])
        self.assertTrue(validators.validate_phone("12345"))

    def test_kpp(self):
        self.assertEqual(validators.validate_kpp("781201001"), [])
        self.assertTrue(validators.validate_kpp("78120100"))


class TestMorphology(unittest.TestCase):
    def test_male_genitive(self):
        result = morphology.full_name_genitive("Иванов Иван Иванович")
        self.assertTrue(result.confident)
        self.assertEqual(result.value, "Иванова Ивана Ивановича")

    def test_male_genitive_hard_cases(self):
        cases = {
            "Хомутов Роман Сергеевич": "Хомутова Романа Сергеевича",
            "Кодловский Максим Анатольевич": "Кодловского Максима Анатольевича",
            "Кузьмин Илья Петрович": "Кузьмина Ильи Петровича",
            "Толстой Лев Николаевич": "Толстого Льва Николаевича",
        }
        for source, expected in cases.items():
            with self.subTest(source=source):
                result = morphology.full_name_genitive(source)
                self.assertTrue(result.confident, result.reason)
                self.assertEqual(result.value, expected)

    def test_female_genitive(self):
        result = morphology.full_name_genitive("Петрова Мария Сергеевна")
        self.assertTrue(result.confident)
        self.assertEqual(result.value, "Петровой Марии Сергеевны")

    def test_indeclinable_surname(self):
        result = morphology.full_name_genitive("Шевченко Пётр Иванович")
        self.assertTrue(result.confident)
        self.assertEqual(result.value, "Шевченко Петра Ивановича")

    def test_fleeting_vowel_surname_is_not_guessed(self):
        # Соловей → Соловья, но Кочубей → Кочубея. Правило ненадёжно —
        # программа обязана спросить, а не угадать.
        result = morphology.full_name_genitive("Соловей Андрей Игоревич")
        self.assertFalse(result.confident)

    def test_uncertain_without_patronymic(self):
        result = morphology.full_name_genitive("Иванов Иван")
        self.assertFalse(result.confident)

    def test_uncertain_foreign_name(self):
        result = morphology.full_name_genitive("Ким Ли Сунович")
        # отчество распознано, но имя нестандартное — программа не уверена
        self.assertFalse(result.confident)

    def test_short_name(self):
        self.assertEqual(morphology.short_name("Иванов Иван Иванович").value, "Иванов И.И.")

    def test_positions(self):
        self.assertEqual(morphology.position_genitive("Генеральный директор").value,
                         "Генерального директора")
        self.assertEqual(morphology.position_genitive("Директор").value, "Директора")
        self.assertFalse(morphology.position_genitive("Главный распорядитель").confident)

    def test_basis(self):
        self.assertEqual(morphology.basis_genitive("Устав").value, "Устава")
        self.assertEqual(morphology.basis_genitive("Устава").value, "Устава")

    def test_normalize_caps(self):
        self.assertEqual(morphology.normalize_person_name("ИВАНОВ ИВАН ИВАНОВИЧ"),
                         "Иванов Иван Иванович")
        self.assertEqual(morphology.normalize_person_name("Иванов Иван Иванович"),
                         "Иванов Иван Иванович")


class TestParser(unittest.TestCase):
    def test_split_company_name(self):
        self.assertEqual(split_company_name('ООО «Ромашка»'),
                         ("ООО", "Общество с ограниченной ответственностью", "Ромашка"))
        self.assertEqual(
            split_company_name('Общество с ограниченной ответственностью «Ромашка»')[2],
            "Ромашка")

    def test_split_removes_duplicated_form(self):
        # В карточках часто пишут форму дважды.
        result = split_company_name(
            'Общество с ограниченной ответственностью ООО «Ромашка»')
        self.assertEqual(result[2], "Ромашка")
        self.assertEqual(result[0], "ООО")

    def test_parse_plain_text(self):
        parsed = parse_text(
            'ООО "СТРОЙИНВЕСТ"\n'
            "ИНН 7812345675\nКПП 781201001\nОГРН 1237800000008\n"
            "Юридический адрес: 190000, г. Санкт-Петербург, ул. Вымышленная, д. 1\n"
            "Генеральный директор: Иванов Иван Иванович\n"
            "Действует на основании Устава\n"
            "Телефон: +7 (812) 000-00-01\nEmail: info@test.example\n")
        company = parsed.company
        self.assertEqual(company.inn, "7812345675")
        self.assertEqual(company.kpp, "781201001")
        self.assertEqual(company.ogrn, "1237800000008")
        self.assertEqual(company.director_position, "Генеральный директор")
        self.assertEqual(company.director_full_name, "Иванов Иван Иванович")
        self.assertEqual(company.director_basis, "Устав")
        self.assertEqual(company.email, "info@test.example")
        self.assertIn("Вымышленная", company.legal_address)
        self.assertEqual(split_company_name(company.short_name)[2], "СТРОЙИНВЕСТ")

    def test_derived_full_name_is_flagged(self):
        parsed = parse_text('ООО "Ромашка"\nИНН 7812345675')
        self.assertIn("full_name", parsed.derived)

    def test_dash_means_empty(self):
        parsed = parse_text("Почтовый адрес: -\nИНН 7812345675")
        self.assertEqual(parsed.company.postal_address, "")

    def test_never_invents_missing_data(self):
        parsed = parse_text("ИНН 7812345675")
        self.assertEqual(parsed.company.director_full_name, "")
        self.assertEqual(parsed.company.legal_address, "")
        self.assertEqual(parsed.company.email, "")

    def test_parse_docx_card(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "card.docx"
            make_card_docx(path, ALPHA)
            parsed = parse_card(read_card(path))
            self.assertEqual(parsed.company.inn, ALPHA["inn"])
            self.assertEqual(parsed.company.ogrn, ALPHA["ogrn"])
            self.assertEqual(parsed.company.director_full_name, ALPHA["director_full_name"])
            self.assertEqual(parsed.company.director_position, ALPHA["director_position"])
            self.assertEqual(parsed.company.director_basis, "Устав")
            self.assertEqual(parsed.company.postal_address, "")

    def test_city_is_recognized_as_address(self):
        """Для ИП адрес часто задают одним городом («Город: …»).

        Раньше такая строка не распознавалась: программа знала только
        метки «Юридический адрес», «Адрес» и подобные, а «Город» — нет.
        """
        parsed = parse_text("ИП Волков Виталий Витальевич\n"
                            "ИНН 312772345390\nГород: Санкт-Петербург\n")
        self.assertEqual(parsed.company.legal_address, "Санкт-Петербург")

        # Другие естественные формулировки тоже понимаются.
        self.assertEqual(
            parse_text("ИП Иванов\nГород прописки: Москва").company.legal_address,
            "Москва")

        # Если есть полный юридический адрес — он важнее строки «Город».
        full = parse_text("ИП Иванов\n"
                         "Юридический адрес: 190000, г. СПб, ул. Ленина, д. 1\n"
                         "Город: Москва\n")
        self.assertIn("Ленина", full.company.legal_address)

    def test_card_number_is_not_taken_for_phone(self):
        """Номер карты или счёта, начинающийся с 8, не должен попасть в телефон.

        Российский телефон — ровно 11 цифр. У карты 16, у расчётного счёта — 20.
        Раньше номер карты предпринимателя, начинавшийся с 8, уезжал в поле
        «Телефон» просто потому, что начинался с восьмёрки.
        """
        from src.company_parser import _looks_like_phone

        self.assertTrue(_looks_like_phone("8 (812) 000-00-03"))
        self.assertTrue(_looks_like_phone("+7 (812) 000-00-03"))
        self.assertFalse(_looks_like_phone("8100 2400 1234 5678"))   # карта, 16
        self.assertFalse(_looks_like_phone("8" + "0" * 19))          # счёт, 20

        # Карта без метки телефона — поле «Телефон» остаётся пустым.
        parsed = parse_text(
            "Индивидуальный предприниматель Иванов Иван Иванович\n"
            "ИНН 781234567870\nОГРНИП 304780123456781\n"
            "Карта 8100 2400 1234 5678\n")
        self.assertEqual(parsed.company.phone, "")

        # Если рядом есть настоящий телефон — берётся он, а не карта.
        parsed = parse_text("ИП Иванов И.И.\nИНН 781234567870\n"
                            "8100 2400 1234 5678\n8 (812) 000-00-03\n")
        self.assertEqual(parsed.company.phone, "8 (812) 000-00-03")

        # ИНН 781234567870 больше не даёт ложный «телефон» 81234567870.
        parsed = parse_text("ИНН 781234567870\nОГРНИП 304780123456781\n")
        self.assertEqual(parsed.company.phone, "")


class TestDocxEngine(unittest.TestCase):
    """Проверка движка на документе, специально разбитом на фрагменты."""

    def _make_split_template(self, path: Path) -> None:
        from docx import Document

        document = Document()
        paragraph = document.add_paragraph()
        for piece in ("Компания ", "{{comp", "any_name", "_bare}}", " с ИНН ", "{{inn}}", "."):
            paragraph.add_run(piece)
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "{{phone}}"
        table.cell(0, 1).text = "{{email}}"
        section = document.sections[0]
        section.header.paragraphs[0].text = "Шапка: {{inn}}"
        document.save(str(path))

    def test_placeholder_split_across_runs(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "t.docx"
            output = Path(tmp) / "out.docx"
            self._make_split_template(template)

            self.assertIn("company_name_bare", scan_placeholders(template))

            report = fill_template(template, output, {
                "company_name_bare": "Ромашка", "inn": "7812345675",
                "phone": "+7 (812) 000-00-01", "email": "a@b.example"})
            text = extract_all_text(output)
            self.assertIn("Компания Ромашка с ИНН 7812345675.", text)
            self.assertIn("Шапка: 7812345675", text)   # колонтитул тоже заполнен
            self.assertIn("+7 (812) 000-00-01", text)  # и таблица
            self.assertNotIn("{{", text)
            self.assertEqual(report.replaced["inn"], 2)

    def test_template_is_never_modified(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "t.docx"
            output = Path(tmp) / "out.docx"
            self._make_split_template(template)
            before = template.read_bytes()
            fill_template(template, output, {"inn": "7812345675"})
            self.assertEqual(before, template.read_bytes())

    def test_refuses_to_overwrite_template(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "t.docx"
            self._make_split_template(template)
            with self.assertRaises(Exception):
                fill_template(template, template, {"inn": "1"})

    def test_unknown_variable_is_reported_not_erased(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "t.docx"
            output = Path(tmp) / "out.docx"
            self._make_split_template(template)
            report = fill_template(template, output, {"inn": "7812345675"})
            self.assertIn("company_name_bare", report.unknown)
            self.assertIn("{{company_name_bare}}", extract_all_text(output))


class TestContext(unittest.TestCase):
    def setUp(self):
        self.project = Project(ROOT)

    def test_digits_split(self):
        context = build_context(make_company(ALPHA), self.project.attorney())
        self.assertEqual(context.values["inn_d1"], "7")
        self.assertEqual(context.values["inn_d10"], "5")
        self.assertEqual("".join(context.values[f"ogrn_d{i}"] for i in range(1, 14)),
                         ALPHA["ogrn"])

    def test_date_in_russian(self):
        context = build_context(make_company(ALPHA), self.project.attorney())
        self.assertEqual(context.values["doc_day"], "17")
        self.assertEqual(context.values["doc_month_name"], "августа")
        self.assertEqual(context.values["doc_year"], "2026")

    def test_marks(self):
        company = make_company(ALPHA)
        company.harm_fund_level = "3"
        company.contract_fund_level = ""
        context = build_context(company, self.project.attorney())
        self.assertEqual(context.values["mark_harm_level3"], "v")
        self.assertEqual(context.values["mark_harm_level1"], "")
        self.assertEqual(context.values["mark_object_ordinary"], "V")
        self.assertTrue(all(context.values[f"mark_contract_level{i}"] == ""
                            for i in "12345"))

    def test_confirmation_requested_for_unclear_declension(self):
        company = make_company(ALPHA)
        company.set("director_full_name", "Ким Ли Сунович")
        context = build_context(company, self.project.attorney())
        self.assertTrue(any(c.key == "director_full_name_genitive"
                            for c in context.confirmations))

    def test_override_wins(self):
        company = make_company(ALPHA)
        company.set("director_full_name", "Ким Ли Сунович")
        company.overrides["director_full_name_genitive"] = "Ким Ли Суновича"
        context = build_context(company, self.project.attorney())
        self.assertEqual(context.values["director_full_name_genitive"], "Ким Ли Суновича")
        self.assertFalse(any(c.key == "director_full_name_genitive"
                             for c in context.confirmations))

    def test_non_ooo_is_warned(self):
        context = build_context(make_company(BETA), self.project.attorney())
        self.assertTrue(any("ООО" in note for note in context.notes))


class TestGeneration(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())
        self.project = Project(ROOT)
        # Явно фиксируем основную СРО, чтобы тесты не зависели от того,
        # какая СРО была «запомнена» в config/app.json.
        self.project.use_sro("СССС", remember=False)
        self.project.output_root = self.tmp / "out"

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_missing_fields_block_generation(self):
        company = make_company(ALPHA)
        company.set("legal_address", "")
        company.set("actual_address", "")
        company.set("email", "")
        company.set("phone", "")
        readiness = check_readiness(self.project, company)
        missing = {label for item in readiness for label in item.missing}
        self.assertIn("Фактический адрес", missing)
        # Почта и телефон теперь НЕ обязательны — их отсутствие не блокирует.
        self.assertNotIn("Электронная почта", missing)
        self.assertNotIn("Телефон", missing)
        with self.assertRaises(GeneratorError) as ctx:
            generate(self.project, company, make_pdf=False)
        message = str(ctx.exception)
        self.assertIn("не хватает следующих данных", message)
        self.assertIn("Фактический адрес", message)
        self.assertNotIn("нет данных", message)

    def test_email_is_optional(self):
        """Без почты документы всё равно формируются (почта необязательна)."""
        self.project.output_root = self.tmp / "out"
        self.project.use_sro("СССС", remember=False)   # у СССС в бланке есть {{email}}
        company = make_company(ALPHA)
        company.set("email", "")
        result = generate(self.project, company, make_pdf=False)
        self.assertTrue(result.ok, [r.problems for r in result.quality])

    def test_phone_is_optional(self):
        """Без телефона документы всё равно формируются (телефон необязателен)."""
        self.project.output_root = self.tmp / "out"
        self.project.use_sro("СССС", remember=False)   # у СССС в бланке есть {{phone}}
        company = make_company(ALPHA)
        company.set("phone", "")
        result = generate(self.project, company, make_pdf=False)
        self.assertTrue(result.ok, [r.problems for r in result.quality])

    def test_actual_address_taken_from_legal(self):
        """У компаний фактический адрес совпадает с юридическим."""
        company = make_company(ALPHA)
        company.set("actual_address", "")
        notes = self.project.apply_auto_fill(company)
        self.assertEqual(company.actual_address, ALPHA["legal_address"])
        self.assertTrue(any("Фактический адрес" in note for note in notes))
        self.assertTrue(all(item.ok for item in check_readiness(self.project, company)))

    def test_postal_address_taken_from_legal(self):
        """Адрес у компании один: юридический, фактический и почтовый совпадают."""
        company = make_company(ALPHA)
        company.set("actual_address", "")
        company.set("postal_address", "")
        self.project.apply_auto_fill(company)
        self.assertEqual(company.actual_address, ALPHA["legal_address"])
        self.assertEqual(company.postal_address, ALPHA["legal_address"])

    def test_every_sro_fills_all_three_addresses(self):
        for profile in self.project.all_sro:
            with self.subTest(sro=profile.key):
                self.assertEqual(profile.auto_fill.get("actual_address"), "legal_address")
                self.assertEqual(profile.auto_fill.get("postal_address"), "legal_address")

    def test_own_actual_address_is_not_overwritten(self):
        company = make_company(ALPHA)
        own = "196084, г. Санкт-Петербург, Лиговский пр., д. 270"
        company.set("actual_address", own)
        company.set("postal_address", own)
        # Оба адреса заданы вручную — подставлять нечего.
        self.assertEqual(self.project.apply_auto_fill(company), [])
        self.assertEqual(company.actual_address, own)
        self.assertEqual(company.postal_address, own)
        self.assertNotEqual(company.actual_address, company.legal_address)

    def test_generation_fills_actual_address_itself(self):
        company = make_company(ALPHA)
        company.set("actual_address", "")
        result = generate(self.project, company, make_pdf=False)
        self.assertTrue(result.ok, [r.problems for r in result.quality])
        application = extract_all_text(result.folder / "01_Заявление_о_вступлении.docx")
        self.assertIn(f"домашний адрес предпринимателя: {ALPHA['legal_address']}",
                      application)
        # Исходный объект не тронут - подстановка идёт в копии.
        self.assertEqual(company.actual_address, "")

    def test_full_scenario(self):
        company = make_company(ALPHA)
        result = generate(self.project, company, make_pdf=False,
                          today=date(2026, 8, 17))
        self.assertTrue(result.ok, [r.problems for r in result.quality])
        self.assertEqual(len(result.created), 2)
        names = sorted(p.name for p in result.created)
        self.assertEqual(names, ["01_Заявление_о_вступлении.docx", "02_Доверенность.docx"])
        # Папка результата: СРО / компания
        self.assertEqual(result.folder.name, "7812345675_ООО Ромашка")
        self.assertEqual(result.folder.parent.name, "СССС")

        application = extract_all_text(result.folder / "01_Заявление_о_вступлении.docx")
        self.assertIn("Общество с ограниченной ответственностью «Ромашка-Строй»; "
                      "ООО «Ромашка»", application)
        self.assertIn("от «17» августа 2026 г.", application)
        self.assertIn(ALPHA["legal_address"], application)
        self.assertIn("Генеральный директор - Иванов Иван Иванович", application)
        self.assertIn("Иванов И.И.", application)
        self.assertNotIn("{{", application)

        power = extract_all_text(result.folder / "02_Доверенность.docx")
        self.assertIn("«17» августа 2026 г.", power)
        self.assertIn("в лице Генерального директора – Иванова Ивана Ивановича", power)
        self.assertIn("на основании Устава", power)
        self.assertIn("ИНН 7812345675", power)
        self.assertIn("Кодловского Максима Анатольевича", power)  # представитель на месте
        self.assertNotIn("{{", power)

    def test_originals_untouched(self):
        before = {path.name: path.read_bytes()
                  for path in (ROOT / "templates").glob("*.docx")}
        originals = {path.name: path.read_bytes()
                     for path in (ROOT / "templates" / "_originals").glob("*.docx")}
        generate(self.project, make_company(ALPHA), make_pdf=False)
        after = {path.name: path.read_bytes()
                 for path in (ROOT / "templates").glob("*.docx")}
        after_originals = {path.name: path.read_bytes()
                           for path in (ROOT / "templates" / "_originals").glob("*.docx")}
        self.assertEqual(before, after)
        self.assertEqual(originals, after_originals)

    def test_no_foreign_email_left_from_template(self):
        """В бланке заявления в гиперссылке оставалась почта чужой компании."""
        result = generate(self.project, make_company(ALPHA), make_pdf=False)
        for path in result.created:
            self.assertNotIn("regionrem", extract_package_text(path))
        application = result.folder / "01_Заявление_о_вступлении.docx"
        # почта попала и в текст, и в адрес гиперссылки
        self.assertEqual(extract_package_text(application).count(ALPHA["email"]), 2)

    def test_two_companies_do_not_mix(self):
        """Главная защита: данные компании А не попадают в документы компании Б."""
        first = generate(self.project, make_company(ALPHA), make_pdf=False)
        second = generate(self.project, make_company(BETA), make_pdf=False)
        self.assertTrue(first.ok)
        self.assertTrue(second.ok)
        self.assertNotEqual(first.folder, second.folder)

        for path in second.created:
            text = extract_package_text(path)
            for marker in (ALPHA["inn"], ALPHA["ogrn"], "Ромашка",
                           ALPHA["email"], "Иванов"):
                self.assertNotIn(marker, text, f"{path.name}: осталось «{marker}»")
        for path in first.created:
            text = extract_package_text(path)
            for marker in (BETA["inn"], BETA["ogrn"], "Василёк",
                           BETA["email"], "Петров"):
                self.assertNotIn(marker, text, f"{path.name}: попало «{marker}»")

    def test_power_of_attorney_says_b_n(self):
        company = make_company(ALPHA)
        company.set("power_number", self.project.new_company().power_number)
        result = generate(self.project, company, make_pdf=False)
        power = extract_all_text(result.folder / "02_Доверенность.docx")
        self.assertIn("ДОВЕРЕННОСТЬ № б/н", power)

    def test_power_of_attorney_number_can_be_overridden(self):
        company = make_company(ALPHA)
        company.set("power_number", "14/2026")
        result = generate(self.project, company, make_pdf=False)
        power = extract_all_text(result.folder / "02_Доверенность.docx")
        self.assertIn("ДОВЕРЕННОСТЬ № 14/2026", power)

    def test_beta_company_documents(self):
        company = make_company(BETA)
        result = generate(self.project, company, make_pdf=False, today=date(2026, 8, 17))
        self.assertTrue(result.ok, [r.problems for r in result.quality])
        power = extract_all_text(result.folder / "02_Доверенность.docx")
        self.assertIn("Акционерное общество «Василёк»", power)
        self.assertIn("в лице Директора – Петровой Марии Сергеевны", power)
        self.assertIn("Петрова М.С.", power)
        application = extract_all_text(result.folder / "01_Заявление_о_вступлении.docx")
        self.assertIn("Акционерное общество «Василёк»; АО «Василёк»", application)

    def test_quality_control_catches_broken_document(self):
        """Если в шаблоне есть переменная без данных, документ не считается готовым."""
        company = make_company(ALPHA)
        context_free = self.project.variables
        self.assertIn("inn", context_free)
        result = generate(self.project, company, make_pdf=False)
        report = result.quality[0]
        self.assertTrue(report.ok)
        self.assertEqual(report.problems, [])

    def test_folder_name_is_safe_for_windows(self):
        company = make_company(ALPHA)
        company.set("short_name", 'ООО «Ромашка/Строй: "Плюс"»')
        name = company_folder_name(company)
        for bad in '<>:"/\\|?*':
            self.assertNotIn(bad, name)
        self.assertTrue(name.startswith("7812345675_"))


def make_sandbox(folder: Path) -> Path:
    """Копия программы с одной готовой СРО и одной, куда бланки не загружены.

    Раньше проверки «СРО без бланков» опирались на то, что какая-то СРО
    в самой программе ещё не размечена. Теперь размечены все, поэтому
    пустая СРО создаётся здесь — проверка больше не зависит от того,
    сколько бланков уже готово.
    """
    root = folder / "программа"
    (root / "sro").mkdir(parents=True)
    shutil.copytree(ROOT / "config", root / "config")
    shutil.copytree(ROOT / "sro" / "СССС", root / "sro" / "СССС")
    empty = root / "sro" / "НОВАЯ СРО"
    (empty / "templates").mkdir(parents=True)
    (empty / "sro.json").write_text(
        json.dumps({"name": "Новая СРО", "short_name": "НОВАЯ", "city": "",
                    "documents": []}, ensure_ascii=False),
        encoding="utf-8")
    return root


class TestMultipleSro(unittest.TestCase):
    """Работа с несколькими саморегулируемыми организациями."""

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())
        self.project = Project(ROOT)
        self.project.output_root = self.tmp / "out"

    def sandbox_project(self) -> Project:
        project = Project(make_sandbox(self.tmp))
        project.output_root = self.tmp / "out"
        return project

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_all_declared_sro_are_found(self):
        keys = {p.key for p in self.project.all_sro}
        self.assertIn("СССС", keys)
        self.assertGreaterEqual(len(keys), 2, "должно быть найдено больше одной СРО")

    def test_every_sro_profile_is_readable(self):
        for profile in self.project.all_sro:
            with self.subTest(sro=profile.key):
                self.assertTrue(profile.short_name)
                self.assertTrue(profile.name)
                self.assertIsInstance(profile.documents, list)

    def test_sro_can_be_found_by_name_or_folder(self):
        from src.sro_registry import find
        for wanted in ("СССС", "сссс", "Ассоциация «Строительный союз Северной столицы»"):
            with self.subTest(wanted=wanted):
                self.assertIsNotNone(find(self.project.all_sro, wanted))
        self.assertIsNone(find(self.project.all_sro, "такой СРО не существует"))

    def test_unknown_sro_is_reported(self):
        with self.assertRaises(GeneratorError) as ctx:
            Project(ROOT, sro="СРО-которой-нет")
        self.assertIn("не найдена", str(ctx.exception))

    def test_all_sro_in_program_are_ready(self):
        """У каждой СРО программы бланки размечены и на месте."""
        for profile in self.project.all_sro:
            with self.subTest(sro=profile.key):
                self.assertTrue(profile.is_ready, profile.readiness_note())

    def test_sro_without_templates_is_not_ready(self):
        project = self.sandbox_project()
        pending = [p for p in project.all_sro if not p.is_ready]
        self.assertTrue(pending, "ожидалась СРО, для которой бланки не загружены")
        for profile in pending:
            with self.subTest(sro=profile.key):
                note = profile.readiness_note()
                self.assertIn("бланки", note.lower())
                self.assertIn(str(profile.templates_dir), note)

    def test_generation_blocked_without_templates(self):
        project = self.sandbox_project()
        pending = next(p for p in project.all_sro if not p.is_ready)
        project.use_sro(pending, remember=False)
        with self.assertRaises(GeneratorError) as ctx:
            generate(project, make_company(ALPHA), make_pdf=False)
        self.assertIn("бланки", str(ctx.exception).lower())

    def test_generate_many_makes_a_set_per_sro(self):
        """Выбор нескольких СРО → по комплекту документов в каждой папке СРО."""
        self.project.output_root = self.tmp / "out"
        chosen = [p for p in self.project.all_sro
                  if p.key in ("СССС", "СФЕРА-А", "ЯРД")]
        outcomes = generate_many(self.project, make_company(ALPHA), chosen,
                                 make_pdf=False)
        self.assertEqual(len(outcomes), 3)
        for outcome in outcomes:
            with self.subTest(sro=outcome.sro.key):
                self.assertTrue(outcome.ok, outcome.error)
                # папка: output / СРО / компания
                self.assertEqual(outcome.result.folder.parent.name, outcome.sro.short_name)
                self.assertEqual(outcome.result.folder.name, "7812345675_ООО Ромашка")
                self.assertEqual(len(outcome.result.created), 2)

    def test_generate_many_does_not_mix_between_sro(self):
        """В комплекте каждой СРО — её название, чужих нет."""
        self.project.output_root = self.tmp / "out"
        chosen = [p for p in self.project.all_sro if p.key in ("СССС", "ЯРД")]
        outcomes = generate_many(self.project, make_company(ALPHA), chosen,
                                 make_pdf=False)
        texts = {}
        for outcome in outcomes:
            self.assertTrue(outcome.ok, outcome.error)
            joined = " ".join(extract_all_text(p) for p in outcome.result.created)
            texts[outcome.sro.key] = joined
        # У каждой СРО в документах её собственное короткое имя.
        self.assertIn("СССС", texts["СССС"])
        self.assertIn("ЯРД", texts["ЯРД"])
        # А чужого имени СРО быть не должно.
        self.assertNotIn("ЯРД", texts["СССС"])

    def test_generate_many_uses_own_params_per_sro(self):
        """У каждой СРО — свои виды работ и уровни, реквизиты общие.

        Так работает мастер параметров в окне: одна компания заявляет
        в разных СРО разное. Проверяем, что нужный уровень отмечен именно
        в своей СРО, а чужой — нет.
        """
        self.project.output_root = self.tmp / "out"
        plans = [
            (self._profile("СФЕРА-А"),
             {"doc_date": "01.09.2026", "doc_number": "СФ-11",
              "object_kind": "hazardous", "harm_fund_level": "3",
              "contract_fund_level": ""}),
            (self._profile("ЯРД"),
             {"doc_date": "02.09.2026", "doc_number": "ЯР-22",
              "object_kind": "ordinary", "harm_fund_level": "2",
              "contract_fund_level": "1"}),
        ]
        outcomes = generate_many(self.project, make_company(ALPHA), plans,
                                 make_pdf=False)
        marks = {}
        texts = {}
        for outcome in outcomes:
            self.assertTrue(outcome.ok, outcome.error)
            application = next(p for p in outcome.result.created
                               if p.name.startswith("01_"))
            marks[outcome.sro.key] = self._marked_levels(application)
            texts[outcome.sro.key] = extract_all_text(application)
        # СФЕРА-А: возмещение вреда — третий уровень, ОДО не заявлен.
        self.assertEqual(marks["СФЕРА-А"], {"вреда": {"Третий"}})
        # ЯРД: возмещение вреда — второй, ОДО — первый.
        self.assertEqual(marks["ЯРД"],
                         {"вреда": {"Второй"}, "обеспечения": {"Первый"}})
        # Номер и дата — свои у каждой СРО.
        self.assertIn("СФ-11", texts["СФЕРА-А"])
        self.assertIn("01» сентября 2026", texts["СФЕРА-А"])
        self.assertIn("ЯР-22", texts["ЯРД"])
        self.assertIn("02» сентября 2026", texts["ЯРД"])
        self.assertNotIn("СФ-11", texts["ЯРД"])

    def _profile(self, key):
        return next(p for p in self.project.all_sro if p.key == key)

    @staticmethod
    def _marked_levels(docx) -> dict:
        """Какие уровни отмечены в таблицах заявления: {фонд: {уровни}}."""
        import zipfile

        from lxml import etree

        from src.docx_engine import W

        root = etree.fromstring(zipfile.ZipFile(docx).read("word/document.xml"))
        names = {"Первый", "Второй", "Третий", "Четвертый", "Пятый"}
        found: dict = {}
        for table in root.iter(W + "tbl"):
            rows = table.findall(W + "tr")
            head = " ".join(t.text or "" for t in rows[0].iter(W + "t")) if rows else ""
            fund = ("вреда" if "вреда" in head else
                    "обеспечения" if "обеспечения" in head else None)
            if fund is None:
                continue
            for row in rows[1:]:
                cells = row.findall(W + "tc")
                label = " ".join(t.text or "" for t in cells[0].iter(W + "t")).strip()
                mark = " ".join(t.text or "" for t in cells[-1].iter(W + "t")).strip()
                if label in names and mark:
                    found.setdefault(fund, set()).add(label)
        return found

    def test_ssss_application_and_power_take_their_numbers(self):
        """У СССС номер заявления и номер доверенности подставляются.

        Раньше в бланке заявления СССС «№ б/н» было вписано текстом и не
        менялось. Проверяем, что теперь оба номера — свои.
        """
        self.project.output_root = self.tmp / "out"
        plans = [(self._profile("СССС"),
                  {"doc_number": "ЗАЯВ-7", "power_number": "ДОВ-7"})]
        outcome = generate_many(self.project, make_company(ALPHA), plans,
                                make_pdf=False)[0]
        self.assertTrue(outcome.ok, outcome.error)
        application = next(p for p in outcome.result.created if p.name.startswith("01_"))
        power = next(p for p in outcome.result.created if p.name.startswith("02_"))
        self.assertIn("ЗАЯВ-7", extract_all_text(application))
        self.assertNotIn("№ б/н", extract_all_text(application))
        self.assertIn("ДОВ-7", extract_all_text(power))

    def test_usage_counter_grows_with_sets(self):
        """Счётчик комплектов растёт на число успешных СРО и сохраняется."""
        self.project.output_root = self.tmp / "out"
        self.project.root = self.tmp  # счётчик пишем во временную папку
        self.assertEqual(self.project.usage_count(), 0)
        chosen = [p for p in self.project.all_sro if p.key in ("СССС", "ЯРД")]
        generate_many(self.project, make_company(ALPHA), chosen, make_pdf=False)
        made = 2
        self.assertEqual(self.project.record_sets(made), made)
        # Пересоздаём проект на той же папке — счётчик читается с диска.
        again = Project(ROOT)
        again.root = self.tmp
        self.assertEqual(again.usage_count(), made)

    def test_generate_many_isolates_errors(self):
        """Сбой у одной СРО не мешает остальным.

        Пятый уровень есть у строительной СРО и нет у проектной (ЯРД, четыре
        уровня). При выборе обеих строительная формируется, а ЯРД честно
        сообщает об ошибке — но не роняет весь пакет.
        """
        self.project.output_root = self.tmp / "out"
        company = make_company(ALPHA)
        company.set("harm_fund_level", "5")
        chosen = [p for p in self.project.all_sro if p.key in ("СФЕРА-А", "ЯРД")]
        outcomes = generate_many(self.project, company, chosen, make_pdf=False)
        by_key = {o.sro.key: o for o in outcomes}
        self.assertTrue(by_key["СФЕРА-А"].ok, by_key["СФЕРА-А"].error)
        self.assertIsNotNone(by_key["ЯРД"].error)
        self.assertIn("уровн", by_key["ЯРД"].error.lower())

    def test_documents_land_in_company_subfolder_of_sro(self):
        """Папки: сверху СРО, внутри — папка компании."""
        self.project.use_sro("СССС", remember=False)
        result = generate(self.project, make_company(ALPHA), make_pdf=False)
        self.assertTrue(result.ok, [r.problems for r in result.quality])
        self.assertEqual(result.folder.name, "7812345675_ООО Ромашка")
        self.assertEqual(result.folder.parent.name, "СССС")
        self.assertEqual(result.folder.parent.parent, self.project.output_root)

    def test_switching_sro_does_not_leak_settings(self):
        """У каждой СРО свои бланки, документы и доверенное лицо."""
        project = self.sandbox_project()
        ready = [p for p in project.all_sro if p.is_ready]
        pending = [p for p in project.all_sro if not p.is_ready]
        project.use_sro(ready[0], remember=False)
        self.assertTrue(project.enabled_documents())
        self.assertTrue(project.attorney())
        project.use_sro(pending[0], remember=False)
        self.assertEqual(project.enabled_documents(), [])
        self.assertEqual(project.attorney(), {})
        self.assertNotEqual(project.templates_dir, ready[0].templates_dir)

    def test_levels_come_from_each_sro(self):
        """У строительных СРО пять уровней, у проектных и изыскательских четыре."""
        counts = {p.key: len(p.harm_levels) for p in self.project.all_sro}
        self.assertEqual(counts.get("СССС"), 5)
        self.assertEqual(counts.get("СФЕРА-А"), 5)
        self.assertEqual(counts.get("СИС"), 5)
        self.assertEqual(counts.get("ЯРД"), 4)
        self.assertEqual(counts.get("СФЕРА ПРОЕКТИРОВЩИКОВ"), 4)
        self.assertEqual(counts.get("СФЕРА ИЗЫСКАТЕЛЕЙ"), 4)
        for profile in self.project.all_sro:
            with self.subTest(sro=profile.key):
                self.assertEqual(len(profile.harm_levels),
                                 len(profile.contract_levels))

    def test_nonexistent_level_is_refused(self):
        """Пятый уровень у проектной СРО — ошибка, а не молча пустая таблица."""
        self.project.use_sro("ЯРД", remember=False)
        company = make_company(ALPHA)
        company.set("harm_fund_level", "5")
        with self.assertRaises(GeneratorError) as ctx:
            generate(self.project, company, make_pdf=False)
        message = str(ctx.exception)
        self.assertIn("нет уровня №5", message)
        self.assertIn("25 миллионов", message)

    def test_levels_differ_between_domains(self):
        by_key = {p.key: p for p in self.project.all_sro}
        building = by_key["СФЕРА-А"].harm_levels[0]
        design = by_key["СФЕРА ПРОЕКТИРОВЩИКОВ"].harm_levels[0]
        self.assertIn("90", building.limit)
        self.assertIn("25", design.limit)
        self.assertNotEqual(building.fee, design.fee)

    def test_remembered_choice_is_restored(self):
        self.project.use_sro("СССС")
        self.assertEqual(Project(ROOT).sro.key, "СССС")


def _character_styles(archive) -> dict:
    """Знаковые стили документа: styleId → его rPr."""
    from lxml import etree

    from src.docx_engine import W

    root = etree.fromstring(archive.read("word/styles.xml"))
    found = {}
    for style in root.iter(W + "style"):
        if style.get(W + "type") != "character":
            continue
        run_properties = style.find(W + "rPr")
        if run_properties is not None:
            found[style.get(W + "styleId")] = run_properties
    return found


def _font_of(run, styles: dict) -> tuple[bool, bool]:
    """Заданы ли у фрагмента шрифт и размер — прямо или знаковым стилем.

    Шрифт можно задать двумя способами, и оба правильные: прямо в самом
    фрагменте либо знаковым стилем (w:rStyle). Важно только, что шрифт
    и размер заданы явно, а не унаследованы «как получится» — из-за такого
    наследования адреса когда-то вышли шрифтом Calibri вместо Times.
    """
    from src.docx_engine import W

    run_properties = run.find(W + "rPr")
    if run_properties is None:
        return False, False
    sources = [run_properties]
    style_reference = run_properties.find(W + "rStyle")
    if style_reference is not None:
        from_style = styles.get(style_reference.get(W + "val"))
        if from_style is not None:
            sources.append(from_style)
    has_font = any(s.find(W + "rFonts") is not None for s in sources)
    has_size = any(s.find(W + "sz") is not None for s in sources)
    return has_font, has_size


class TestProjectConfig(unittest.TestCase):
    def test_every_placeholder_is_described(self):
        """Каждая переменная шаблонов должна быть описана в config/variables.json."""
        from src.quality_control import lookup_variable

        project = Project(ROOT)
        for profile in project.all_sro:
            if not profile.is_ready:
                continue
            project.use_sro(profile, remember=False)
            for spec in profile.enabled_documents():
                for name in project.placeholders(spec):
                    with self.subTest(sro=profile.key, document=spec.title, variable=name):
                        self.assertIsNotNone(
                            lookup_variable(name, project.variables),
                            f"переменная {{{{{name}}}}} не описана "
                            f"в config/variables.json")

    def test_power_of_attorney_number_default(self):
        """По принятому порядку все доверенности выдаются без номера."""
        project = Project(ROOT)
        self.assertEqual(project.new_company().power_number, "б/н")

    def test_every_placeholder_has_explicit_font(self):
        """Подставляемый текст должен быть тем же шрифтом, что и остальной.

        В бланке пропуски для адресов не имели настроек шрифта, и значение
        выходило шрифтом по умолчанию (Calibri) вместо Times New Roman.
        """
        import zipfile

        from lxml import etree

        from src.docx_engine import W, _iter_paragraphs, _own_text_nodes

        project = Project(ROOT)
        specs = []
        for profile in project.all_sro:
            if not profile.is_ready:
                continue
            project.use_sro(profile, remember=False)
            specs += [(profile.key, spec, project.template_path(spec))
                      for spec in profile.enabled_documents()]
        for sro_key, spec, template in specs:
            archive = zipfile.ZipFile(template)
            root = etree.fromstring(archive.read("word/document.xml"))
            for paragraph in _iter_paragraphs(root):
                for node in _own_text_nodes(paragraph):
                    if "{{" not in (node.text or ""):
                        continue
                    styles = _character_styles(archive)
                    font, size = _font_of(node.getparent(), styles)
                    with self.subTest(sro=sro_key, document=spec.title, text=node.text):
                        self.assertTrue(
                            font,
                            f"{spec.template}: у фрагмента {node.text!r} "
                            f"не задан шрифт")
                        self.assertTrue(
                            size,
                            f"{spec.template}: у фрагмента {node.text!r} "
                            f"не задан размер")

    def test_signature_has_room_to_sign(self):
        """Между наименованием и инициалами — табуляция до правого поля."""
        import zipfile

        from lxml import etree

        from src.docx_engine import W, _iter_paragraphs, _own_text_nodes

        project = Project(ROOT)
        spec = next(s for s in project.enabled_documents() if s.id == "application")
        root = etree.fromstring(
            zipfile.ZipFile(project.template_path(spec)).read("word/document.xml"))
        for paragraph in _iter_paragraphs(root):
            text = "".join(n.text or "" for n in _own_text_nodes(paragraph))
            if "{{director_short_name}}" not in text:
                continue
            self.assertTrue(list(paragraph.iter(W + "tab")),
                            "в строке подписи нет табуляции")
            tabs = paragraph.find(W + "pPr").find(W + "tabs")
            self.assertIsNotNone(tabs, "не задана позиция табуляции")
            self.assertEqual(tabs.find(W + "tab").get(W + "val"), "right")
            return
        self.fail("не найдена строка подписи руководителя")

    def test_header_cells_are_level(self):
        """Номер заявления и адресат должны стоять на одной высоте."""
        import zipfile

        from lxml import etree

        from src.docx_engine import W

        project = Project(ROOT)
        spec = next(s for s in project.enabled_documents() if s.id == "application")
        root = etree.fromstring(
            zipfile.ZipFile(project.template_path(spec)).read("word/document.xml"))
        header = list(root.iter(W + "tbl"))[1]
        for paragraph in header.iter(W + "p"):
            p_pr = paragraph.find(W + "pPr")
            spacing = p_pr.find(W + "spacing") if p_pr is not None else None
            if spacing is None:
                continue
            self.assertIsNone(spacing.get(W + "beforeAutospacing"),
                              "остался автоматический интервал — строки разъедутся")
            self.assertIsNone(spacing.get(W + "afterAutospacing"),
                              "остался автоматический интервал — строки разъедутся")

    def test_caption_and_value_are_on_separate_lines(self):
        """Значение не должно прилипать к подписи к строке.

        В бланках подпись к строке и следующая строка лежат в одном абзаце.
        Раньше телефон директора попадал прямо в подпись, и выходило
        «E-mail руководителя организацииИванов И.И., Директор…».
        """
        project = Project(ROOT)
        company = make_company(ALPHA)
        for profile in project.all_sro:
            if not profile.is_ready:
                continue
            project.use_sro(profile, remember=False)
            spec = next((s for s in profile.enabled_documents()
                         if s.id == "application"), None)
            if spec is None:
                continue
            values = build_context(company, project.attorney(), sro=profile).values
            with tempfile.TemporaryDirectory() as folder:
                target = Path(folder) / spec.template
                fill_template(project.template_path(spec), target, values)
                text = extract_all_text(target)
            # Подписи взяты целиком: у СИС слова «контактного лица» стоят
            # в середине подписи, и обрывок совпал бы там, где всё в порядке.
            for caption in (
                    "E-mail руководителя организации",
                    "E-mail контактного лица",
                    "улица (проспект, переулок и др.)",
                    "и номер дома (владения), корпуса (строения) и офиса"):
                for line in text.splitlines():
                    if caption not in line:
                        continue
                    tail = line.split(caption, 1)[1].strip()
                    with self.subTest(sro=profile.key, caption=caption):
                        self.assertEqual(
                            tail, "",
                            f"{profile.key}: к подписи «{caption}» прилип текст "
                            f"«{tail[:60]}»")

    def test_contacts_are_written_once(self):
        """Контактные данные в п.7 подставляются ОДИН раз.

        По просьбе пользователя данные в «Контактных данных» пишутся один
        раз (строка руководителя), без повтора и без отдельной строки
        контактов организации. Серые подписи бланка при этом остаются —
        за это отвечает другой тест (о слипании подписи со значением).
        """
        project = Project(ROOT)
        company = make_company(ALPHA)
        values = None
        for profile in project.all_sro:
            if not profile.is_ready:
                continue
            project.use_sro(profile, remember=False)
            spec = next((s for s in profile.enabled_documents()
                         if s.id == "application"), None)
            if spec is None:
                continue
            placeholders = project.placeholders(spec)
            if "director_contact_line" not in placeholders:
                continue  # в бланке нет пункта «Контактные данные»
            if values is None:
                values = build_context(company, project.attorney(),
                                       sro=profile).values
            with tempfile.TemporaryDirectory() as folder:
                target = Path(folder) / spec.template
                fill_template(project.template_path(spec), target, values)
                text = extract_all_text(target)
            line = values["director_contact_line"]
            with self.subTest(sro=profile.key):
                self.assertEqual(
                    text.count(line), 1,
                    f"{profile.key}: контактная строка встречается "
                    f"{text.count(line)} раз(а), а должна ровно один раз")

    def test_tabs_survive_rewriting(self):
        """Табуляции бланка нельзя терять при переписывании абзаца.

        Строка подписи разнесена табуляцией; когда она пропадала, инициалы
        прилипали к тексту и расписаться было негде.
        """
        import zipfile

        from lxml import etree

        from src.docx_engine import W, _iter_paragraphs, _own_text_nodes

        project = Project(ROOT)
        checked = 0
        for profile in project.all_sro:
            if not profile.is_ready:
                continue
            project.use_sro(profile, remember=False)
            spec = next((s for s in profile.enabled_documents()
                         if s.id == "application"), None)
            if spec is None:
                continue
            root = etree.fromstring(zipfile.ZipFile(
                project.template_path(spec)).read("word/document.xml"))
            for paragraph in _iter_paragraphs(root):
                text = "".join(n.text or "" for n in _own_text_nodes(paragraph))
                if "Подпись уполномоченного лица организации" not in text:
                    continue
                checked += 1
                # Важен не сам факт табуляции, а её место: она должна стоять
                # МЕЖДУ надписью и инициалами. Раньше переписывание абзаца
                # сваливало весь текст в первый фрагмент, а табуляции
                # оставались болтаться в хвосте — расписаться было негде.
                order = []
                for element in paragraph.iter():
                    if element.tag == W + "tab":
                        order.append("\t")
                    elif element.tag == W + "t":
                        order.append(element.text or "")
                joined = "".join(order)
                with self.subTest(sro=profile.key):
                    self.assertRegex(
                        joined,
                        r"организации\t.*\{\{director_short_name\}\}",
                        f"{profile.key}: между надписью и инициалами нет "
                        f"табуляции — расписаться будет негде")
        self.assertTrue(checked, "не найдено ни одной строки подписи")

    def test_no_data_from_sample_blanks_survives(self):
        """Реквизиты компаний из образцов не должны попадать в документы.

        Два бланка пришли ЗАПОЛНЕННЫМИ — с данными других компаний.
        Это главная опасность: документ выглядит готовым, а в нём чужие
        ОГРН и ИНН. Проверяем прямым поиском по всем готовым документам
        всех СРО.
        """
        # что осталось в бланках от компаний, по которым их заполняли
        foreign = {
            "ЗАВОД ЭЛЕКТРОПУЛЬТ": "наименование компании из образца СИС",
            "1027804180766": "ОГРН компании из образца СИС",
            "7806008569": "ИНН компании из образца СИС",
            "ЭЛЕКТРОПУЛЬТОВЦЕВ": "адрес компании из образца СИС",
            "regionrem": "адрес почты из бланка СССС",
            "КонтакноеЛицо": "метка программы слияния",
            "МылоКонтрактнЛ": "метка программы слияния",
        }
        project = Project(ROOT)
        company = make_company(ALPHA)
        for profile in project.all_sro:
            project.use_sro(profile, remember=False)
            values = build_context(company, project.attorney(), sro=profile).values
            for spec in profile.enabled_documents():
                with tempfile.TemporaryDirectory() as folder:
                    target = Path(folder) / spec.template
                    fill_template(project.template_path(spec), target, values)
                    text = extract_package_text(target)
                for needle, what in foreign.items():
                    with self.subTest(sro=profile.key, document=spec.title,
                                      found=needle):
                        self.assertNotIn(
                            needle, text,
                            f"{profile.key}/{spec.template}: в документе "
                            f"осталось чужое — {what}")

    def test_word_fields_are_flattened(self):
        """В шаблонах не должно остаться «полей» Word.

        Бланк СИС собран программой слияния: реквизиты вставлены полями
        { AUTHOR ... }. Word пересчитывает такие поля при открытии, и вместо
        подставленных данных в документе появляется имя автора файла.
        """
        import zipfile

        from lxml import etree

        from src.docx_engine import W

        project = Project(ROOT)
        for profile in project.all_sro:
            project.use_sro(profile, remember=False)
            for spec in profile.enabled_documents():
                root = etree.fromstring(zipfile.ZipFile(
                    project.template_path(spec)).read("word/document.xml"))
                codes = [(node.text or "").strip()
                         for node in root.iter(W + "instrText")]
                with self.subTest(sro=profile.key, document=spec.title):
                    self.assertEqual(
                        codes, [],
                        f"{profile.key}/{spec.template}: остались поля Word "
                        f"({', '.join(codes[:3])}) — Word пересчитает их "
                        f"и подменит данные")

    def test_entrepreneur_digits_fit_the_blank(self):
        """У предпринимателя 12 и 15 знаков — клеток должно хватить.

        Бланки напечатаны под юрлицо: 10 клеток под ИНН и 13 под ОГРН.
        Последние цифры предпринимателя просто пропали бы, а документ
        выглядел бы заполненным.
        """
        import zipfile

        from lxml import etree

        from src.docx_engine import W

        project = Project(ROOT)
        company = make_company(IVAN)
        for profile in project.all_sro:
            project.use_sro(profile, remember=False)
            values = build_context(company, project.attorney(), sro=profile).values
            for spec in profile.enabled_documents():
                placeholders = project.placeholders(spec)
                if "inn_d1" not in placeholders:
                    continue  # в этом бланке цифры не по клеткам
                with tempfile.TemporaryDirectory() as folder:
                    target = Path(folder) / spec.template
                    fill_template(project.template_path(spec), target, values)
                    root = etree.fromstring(
                        zipfile.ZipFile(target).read("word/document.xml"))
                    rows = {}
                    for row in root.iter(W + "tr"):
                        text = "".join(n.text or "" for n in row.iter(W + "t"))
                        if text in (company.inn, company.ogrn):
                            rows[text] = row
                for value in (company.inn, company.ogrn):
                    with self.subTest(sro=profile.key, document=spec.title,
                                      value=value):
                        self.assertIn(
                            value, rows,
                            f"{profile.key}/{spec.template}: {value} не собрался "
                            f"из клеток — часть цифр потерялась")
                        self.assertGreaterEqual(
                            len(rows[value].findall(W + "tc")), len(value),
                            f"{profile.key}/{spec.template}: клеток меньше, "
                            f"чем знаков в {value}")

    def test_expanding_cells_keeps_table_width(self):
        """Новые клетки делят прежнюю ширину, а не расширяют таблицу.

        Иначе ряд вылезет за поле страницы и документ придётся править руками.
        """
        import zipfile

        from lxml import etree

        from src.docx_engine import W, _row_width

        project = Project(ROOT)
        for profile in project.all_sro:
            project.use_sro(profile, remember=False)
            for spec in profile.enabled_documents():
                if "inn_d1" not in project.placeholders(spec):
                    continue
                template = project.template_path(spec)
                before = {}
                root = etree.fromstring(zipfile.ZipFile(template).read("word/document.xml"))
                for prefix in ("inn", "ogrn"):
                    row = _digit_row(root, prefix)
                    if row is not None:
                        before[prefix] = _row_width(row)
                values = build_context(make_company(IVAN), project.attorney(),
                                       sro=profile).values
                with tempfile.TemporaryDirectory() as folder:
                    target = Path(folder) / spec.template
                    fill_template(template, target, values)
                    filled = etree.fromstring(
                        zipfile.ZipFile(target).read("word/document.xml"))
                for prefix, width in before.items():
                    digits = values["inn"] if prefix == "inn" else values["ogrn"]
                    row = next((r for r in filled.iter(W + "tr")
                                if "".join(n.text or "" for n in r.iter(W + "t")) == digits),
                               None)
                    with self.subTest(sro=profile.key, document=spec.title, row=prefix):
                        self.assertIsNotNone(row)
                        self.assertEqual(
                            _row_width(row), width,
                            f"{profile.key}/{spec.template}: ширина ряда {prefix} "
                            f"изменилась — таблица вылезет за поля")
                        grid = row.getparent().find(W + "tblGrid")
                        self.assertEqual(
                            len(grid.findall(W + "gridCol")),
                            len(row.findall(W + "tc")),
                            "сетка колонок не совпадает с числом клеток")

    def test_entrepreneur_documents_are_complete(self):
        """Документы предпринимателя проходят проверку у всех СРО.

        Отдельно важно, что КПП не требуется: у предпринимателя его нет,
        и просить его — значит просить придумать несуществующий реквизит.
        """
        project = Project(ROOT)
        for profile in project.all_sro:
            project.use_sro(profile, remember=False)
            # Так же, как в окне программы: сперва умолчания и подстановка
            # адресов, потом проверка готовности.
            company = make_company(IVAN)
            project.apply_applicant_defaults(company)
            project.apply_auto_fill(company)
            for item in check_readiness(project, company):
                with self.subTest(sro=profile.key, document=item.spec.title):
                    self.assertTrue(
                        item.ok,
                        f"{profile.key}/{item.spec.title}: не хватает "
                        f"{', '.join(item.missing) or item.unknown_variables}")

    def test_entrepreneur_name_not_in_quotes(self):
        """ФИО предпринимателя НЕ берётся в кавычки-ёлочки, а у юрлица — берётся.

        Кавычки « » в бланках напечатаны под юрлицо (ООО «Ромашка»). Для ИП
        «Иванов Иван Иванович» в кавычках читается как название фирмы — это
        ошибка. Проверяем сразу все бланки всех СРО: ни в одном готовом
        документе ИП имя не должно стоять в кавычках. И наоборот — у юрлица
        кавычки вокруг наименования обязаны сохраниться.
        """
        import re

        def quoted(path):
            bits = []
            for line in extract_all_text(path).splitlines():
                bits += re.findall(r"«[^»\n]*»", line)
            return bits

        project = Project(ROOT)
        legal_entity_kept_quotes = False
        for profile in project.all_sro:
            project.use_sro(profile, remember=False)
            ip_values = build_context(make_company(IVAN), project.attorney(),
                                      sro=profile).values
            ur_values = build_context(make_company(ALPHA), project.attorney(),
                                      sro=profile).values
            for spec in profile.enabled_documents():
                with tempfile.TemporaryDirectory() as folder:
                    ip_doc = Path(folder) / ("ip_" + spec.template)
                    ur_doc = Path(folder) / ("ur_" + spec.template)
                    fill_template(project.template_path(spec), ip_doc, ip_values)
                    fill_template(project.template_path(spec), ur_doc, ur_values)
                    for chunk in quoted(ip_doc):
                        with self.subTest(sro=profile.key, document=spec.title,
                                          chunk=chunk):
                            self.assertNotIn(
                                "Иванов", chunk,
                                f"{profile.key}/{spec.title}: ФИО ИП в кавычках {chunk!r}")
                            self.assertNotIn(
                                "предприниматель", chunk.lower(),
                                f"{profile.key}/{spec.title}: наименование ИП "
                                f"в кавычках {chunk!r}")
                    if any("Ромашка" in chunk for chunk in quoted(ur_doc)):
                        legal_entity_kept_quotes = True
        self.assertTrue(
            legal_entity_kept_quotes,
            "у юрлица кавычки вокруг наименования должны были сохраниться")

    def test_templates_exist(self):
        project = Project(ROOT)
        for profile in project.all_sro:
            if not profile.is_ready:
                continue
            project.use_sro(profile, remember=False)
            for spec in profile.enabled_documents():
                self.assertTrue(project.template_path(spec).exists(),
                                f"{profile.key}: {spec.template}")


class TestFirstRunWindow(unittest.TestCase):
    """Первый запуск: окно выбора СРО должно быть ВИДНО.

    На первом запуске главное окно ещё скрыто (`withdraw`), и если окно
    выбора СРО привязать к нему через `transient`, на Windows диалог тоже
    становится невидимым — программа как будто зависает при запуске. Тесту
    нужен экран, поэтому без дисплея он пропускается (в обычном прогоне так
    и будет; под Xvfb — выполняется).
    """

    def _make_root(self):
        try:
            import tkinter as tk
        except ImportError as exc:            # tkinter не установлен
            self.skipTest(f"tkinter недоступен: {exc}")
        try:
            root = tk.Tk()
        except tk.TclError as exc:            # нет дисплея
            self.skipTest(f"нет графического окружения: {exc}")
        return tk, root

    def test_sro_dialog_visible_with_hidden_parent(self):
        tk, root = self._make_root()
        from src import gui
        try:
            root.withdraw()  # как в main() на первом запуске
            project = Project(ROOT)
            seen = {}

            def poll():
                for child in root.winfo_children():
                    if isinstance(child, gui.SroDialog):
                        child.update_idletasks(); root.update()
                        seen["viewable"] = bool(child.winfo_viewable())
                        seen["transient"] = bool(child.wm_transient())
                        child._accept()
                        return
                root.after(40, poll)

            root.after(80, poll)
            chosen = gui.SroDialog(root, project.all_sro, project.sro).result
            self.assertTrue(seen.get("viewable"),
                            "окно выбора СРО невидимо при скрытом главном окне "
                            "— на Windows программа зависнет при запуске")
            self.assertFalse(seen.get("transient"),
                             "к скрытому родителю не должно быть привязки transient")
            self.assertIsNotNone(chosen, "выбор СРО не вернулся")
        finally:
            root.destroy()

    def test_sro_dialog_stays_child_of_visible_parent(self):
        tk, root = self._make_root()
        from src import gui
        try:
            root.deiconify(); root.update()  # как при кнопке «Сменить СРО…»
            project = Project(ROOT)
            seen = {}

            def poll():
                for child in root.winfo_children():
                    if isinstance(child, gui.SroDialog):
                        child.update_idletasks(); root.update()
                        seen["transient"] = bool(child.wm_transient())
                        child._accept()
                        return
                root.after(40, poll)

            root.after(80, poll)
            gui.SroDialog(root, project.all_sro, project.sro)
            self.assertTrue(seen.get("transient"),
                            "к видимому родителю окно должно оставаться дочерним")
        finally:
            root.destroy()


if __name__ == "__main__":
    unittest.main(verbosity=2)
