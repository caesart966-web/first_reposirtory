# -*- coding: utf-8 -*-
"""Чтение карточек компании из файлов.

Все читалки возвращают одно и то же: `CardContent` — плоский текст плюс
пары «подпись → значение», найденные в таблицах. Разбором смысла занимается
company_parser.py, а не эти функции.

Обработка полностью локальная. Ничего никуда не отправляется.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

SUPPORTED_SUFFIXES = {".doc", ".docx", ".xlsx", ".xlsm", ".pdf", ".txt", ".csv",
                      ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}


class ReadError(Exception):
    """Ошибка чтения файла с текстом, понятным пользователю."""


@dataclass
class CardContent:
    """Содержимое карточки компании в сыром виде."""

    text: str = ""
    pairs: list[tuple[str, str]] = field(default_factory=list)
    source: str = ""

    def merged_text(self) -> str:
        """Текст + пары в виде «подпись: значение» — так парсеру удобнее."""
        lines = [f"{k}: {v}" for k, v in self.pairs if v.strip()]
        return "\n".join(lines + [self.text])


# ---------------------------------------------------------------- DOCX
def read_docx(path: Path) -> CardContent:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as exc:  # pragma: no cover
        raise ReadError("Не установлена библиотека python-docx для чтения DOCX.") from exc

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ReadError(
            f"Не удалось открыть файл «{path.name}». Возможно, он повреждён "
            f"или это не документ Word."
        ) from exc

    lines: list[str] = []
    pairs: list[tuple[str, str]] = []

    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())

    def harvest_table(table) -> None:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # Схлопываем объединённые ячейки, дающие дубли подряд.
            deduped: list[str] = []
            for cell in cells:
                if not deduped or deduped[-1] != cell:
                    deduped.append(cell)
            if len(deduped) == 2 and deduped[0]:
                pairs.append((deduped[0], deduped[1]))
            else:
                joined = " | ".join(c for c in deduped if c)
                if joined:
                    lines.append(joined)
            for cell in row.cells:
                for nested in cell.tables:
                    harvest_table(nested)

    for table in doc.tables:
        harvest_table(table)

    # Колонтитулы тоже иногда несут реквизиты.
    for section in doc.sections:
        for container in (section.header, section.footer):
            for para in container.paragraphs:
                if para.text.strip():
                    lines.append(para.text.strip())

    return CardContent("\n".join(lines), pairs, path.name)


# ---------------------------------------------------------------- XLSX
def read_xlsx(path: Path) -> CardContent:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise ReadError("Не установлена библиотека openpyxl для чтения Excel-файлов.") from exc

    try:
        book = load_workbook(str(path), data_only=True, read_only=True)
    except Exception as exc:
        raise ReadError(
            f"Не удалось открыть файл «{path.name}». Возможно, он повреждён или "
            f"это файл старого формата .xls — пересохраните его в Excel как .xlsx."
        ) from exc

    lines: list[str] = []
    pairs: list[tuple[str, str]] = []
    try:
        for sheet in book.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if v is None else str(v).strip() for v in row]
                cells = [c for c in cells if c]
                if not cells:
                    continue
                if len(cells) == 2:
                    pairs.append((cells[0], cells[1]))
                else:
                    lines.append(" | ".join(cells))
    finally:
        book.close()

    return CardContent("\n".join(lines), pairs, path.name)


# ---------------------------------------------------------------- PDF
def read_pdf(path: Path) -> CardContent:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ReadError("Не установлена библиотека pypdf для чтения PDF.") from exc

    try:
        reader = PdfReader(str(path))
        chunks = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise ReadError(
            f"Не удалось прочитать PDF «{path.name}». Возможно, файл повреждён "
            f"или защищён паролем."
        ) from exc

    text = "\n".join(chunks).strip()
    if not text:
        raise ReadError(
            f"В PDF «{path.name}» не найдено ни одной строки текста. "
            f"Скорее всего, это скан (картинка). Скопируйте реквизиты вручную "
            f"во вкладку «Вставить текст»."
        )
    return CardContent(text, [], path.name)


# ---------------------------------------------------------------- изображение
def read_image(path: Path) -> CardContent:
    """Распознавание текста с картинки. Требует локально установленный Tesseract."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        raise ReadError(
            "Распознавание текста с изображений не настроено.\n\n"
            "Что делать: откройте картинку, скопируйте реквизиты и вставьте их "
            "во вкладку «Вставить текст» — это быстрее и надёжнее.\n\n"
            "Если распознавание нужно постоянно, установите программу Tesseract OCR "
            "и библиотеки: pip install pytesseract pillow"
        ) from exc

    try:
        text = pytesseract.image_to_string(Image.open(str(path)), lang="rus+eng")
    except Exception as exc:
        raise ReadError(
            f"Не удалось распознать текст на изображении «{path.name}». "
            f"Проверьте, что установлена программа Tesseract OCR с русским языком. "
            f"Техническая причина: {exc}"
        ) from exc

    if not text.strip():
        raise ReadError(
            f"На изображении «{path.name}» текст не распознан. "
            f"Попробуйте более чёткий снимок или вставьте реквизиты текстом."
        )
    return CardContent(text, [], path.name)


# ---------------------------------------------------------------- текст
def read_text_file(path: Path) -> CardContent:
    for encoding in ("utf-8-sig", "cp1251", "utf-16"):
        try:
            return CardContent(path.read_text(encoding=encoding), [], path.name)
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ReadError(
        f"Не удалось определить кодировку файла «{path.name}». "
        f"Сохраните его в формате UTF-8 или вставьте текст вручную."
    )


# ---------------------------------------------------------------- старый .doc
def _convert_doc_with_libreoffice(source: Path, folder: Path) -> Path | None:
    """Преобразовать .doc в .docx установленным LibreOffice."""
    import subprocess

    from .pdf_export import find_libreoffice

    executable = find_libreoffice()
    if not executable:
        return None
    try:
        subprocess.run(
            [executable, "--headless", "--norestore", "--convert-to", "docx",
             "--outdir", str(folder), str(source)],
            check=True, capture_output=True, timeout=180)
    except Exception:      # noqa: BLE001 — не вышло, попробуем Word
        return None
    target = folder / (source.stem + ".docx")
    return target if target.exists() else None


def _convert_doc_with_word(source: Path, folder: Path) -> Path | None:
    """Преобразовать .doc в .docx установленным Microsoft Word (Windows)."""
    try:
        import pythoncom
        import win32com.client
    except ImportError:
        return None

    target = folder / (source.stem + ".docx")
    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        document = word.Documents.Open(
            str(source), ConfirmConversions=False, ReadOnly=True,
            AddToRecentFiles=False)
        try:
            document.SaveAs2(str(target), FileFormat=16)   # 16 — обычный .docx
        finally:
            document.Close(False)
    except Exception:      # noqa: BLE001 — Word не установлен или занят
        return None
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:      # noqa: BLE001
                pass
        pythoncom.CoUninitialize()
    return target if target.exists() else None


def read_doc(path: Path) -> CardContent:
    """Прочитать карточку в старом формате Word (.doc).

    Такой файл — не ZIP с XML, как .docx, а двоичный формат, читать его
    напрямую нечем. Поэтому программа сама пересохраняет его в .docx:
    сначала пробует LibreOffice (он не показывает окон и не зависает —
    стоит ограничение по времени), затем установленный Microsoft Word.
    Исходный файл при этом не меняется: работа идёт с временной копией.
    """
    import tempfile

    with tempfile.TemporaryDirectory() as folder:
        folder = Path(folder)
        # Копию кладём рядом: конвертеры пишут результат по имени исходника,
        # а исходную папку пользователя трогать нельзя.
        copy = folder / path.name
        try:
            copy.write_bytes(path.read_bytes())
        except OSError as exc:
            raise ReadError(f"Не удалось прочитать файл «{path.name}».") from exc

        converted = (_convert_doc_with_libreoffice(copy, folder)
                     or _convert_doc_with_word(copy, folder))
        if converted is None:
            raise ReadError(
                f"Файл «{path.name}» — старого формата Word (.doc). "
                f"Преобразовать его не удалось: на компьютере не найден "
                f"ни Microsoft Word, ни LibreOffice.\n\n"
                f"Откройте файл в Word и сохраните как .docx "
                f"(Файл — Сохранить как — тип «Документ Word»), "
                f"затем загрузите снова."
            )
        return read_docx(converted)


# ---------------------------------------------------------------- точка входа
READERS = {
    ".doc": read_doc,
    ".docx": read_docx,
    ".xlsx": read_xlsx,
    ".xlsm": read_xlsx,
    ".pdf": read_pdf,
    ".txt": read_text_file,
    ".csv": read_text_file,
    ".png": read_image,
    ".jpg": read_image,
    ".jpeg": read_image,
    ".tif": read_image,
    ".tiff": read_image,
    ".bmp": read_image,
}


def read_card(path: str | Path) -> CardContent:
    """Прочитать карточку компании из файла любого поддерживаемого формата."""
    path = Path(path)
    if not path.exists():
        raise ReadError(f"Файл не найден: {path}")
    if path.is_dir():
        raise ReadError(f"«{path.name}» — это папка, а не файл карточки.")

    suffix = path.suffix.lower()
    if suffix not in READERS:
        supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
        raise ReadError(
            f"Формат файла «{suffix or 'без расширения'}» не поддерживается.\n"
            f"Поддерживаются: {supported}"
        )
    return READERS[suffix](path)


def read_pasted_text(text: str) -> CardContent:
    """Карточка из текста, вставленного в окно программы."""
    if not (text or "").strip():
        raise ReadError("Текст пустой — вставьте реквизиты компании.")
    return CardContent(_normalize_whitespace(text), [], "вставленный текст")


def _normalize_whitespace(text: str) -> str:
    text = text.replace(" ", " ").replace("﻿", "")
    return re.sub(r"[ \t]+", " ", text)
